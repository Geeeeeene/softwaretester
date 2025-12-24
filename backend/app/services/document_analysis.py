"""文档分析服务 - 读取docx文件并使用AI总结要点"""
import os
import sys
import logging
from typing import Optional, Dict, Any
from pathlib import Path
import anthropic

from app.core.config import settings

logger = logging.getLogger(__name__)

class DocumentAnalysisService:
    """文档分析服务 - 分析设计文档并总结要点"""
    
    def __init__(self, api_key: Optional[str] = None, model: Optional[str] = None, base_url: Optional[str] = None):
        self.api_key = api_key or settings.CLAUDE_API_KEY
        self.model = model or settings.CLAUDE_MODEL
        self.base_url = base_url or settings.CLAUDE_BASE_URL
        
        if not self.api_key:
            logger.warning("CLAUDE_API_KEY 未设置，文档分析功能将不可用")
    
    def read_docx(self, file_path: Path) -> str:
        """读取docx文件内容"""
        try:
            from docx import Document
            doc = Document(file_path)
            
            # 提取所有段落文本
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            content = "\n".join(paragraphs)
            print(f"📄 读取docx文件成功，长度: {len(content)} 字符", file=sys.stderr, flush=True)
            return content
            
        except ImportError:
            raise Exception("python-docx 库未安装，请运行: pip install python-docx")
        except Exception as e:
            logger.error(f"读取docx文件失败: {str(e)}")
            raise Exception(f"读取docx文件失败: {str(e)}")
    
    async def analyze_document(self, docx_path: Path) -> str:
        """使用AI分析文档并总结要点"""
        try:
            # 读取文档内容
            doc_content = self.read_docx(docx_path)
            
            if not doc_content or len(doc_content.strip()) < 50:
                raise Exception("文档内容过少，无法进行分析")
            
            # 构建分析提示词
            # 限制文档长度避免token过多
            doc_content_limited = doc_content[:8000] if len(doc_content) > 8000 else doc_content
            prompt = f"""你是一个专业的技术文档分析师。请仔细阅读以下设计文档，并总结出关键要点。

**重要说明**：这些要点将用于帮助AI生成单元测试用例，因此需要重点关注对测试用例生成有价值的信息。

**文档内容**:
{doc_content_limited}

**要求**:
1. 总结文档的核心设计理念和架构思路（有助于理解代码的整体结构和设计意图）
2. 提取关键的功能模块和组件说明（帮助理解各模块的职责和功能）
3. 总结重要的接口定义和使用规范（测试用例需要验证接口的正确性）
4. 提取关键的约束条件和注意事项（测试用例需要覆盖边界情况和异常情况）
5. 特别关注与测试相关的信息：
   - 函数的预期行为和返回值
   - 错误处理机制
   - 边界条件和特殊场景
   - 依赖关系和前置条件
6. 总结要点要简洁明了，控制在500字以内
7. 使用中文输出

**目标**：生成的要点应该能够帮助测试用例生成AI更好地理解代码的设计意图、功能要求和测试重点，从而生成更准确、更全面的单元测试用例。

请生成文档要点总结："""

            if self.base_url:
                client = anthropic.Anthropic(api_key=self.api_key, base_url=self.base_url)
            else:
                client = anthropic.Anthropic(api_key=self.api_key)
            
            print(f"📤 正在发送请求到 Claude API 分析文档...", file=sys.stderr, flush=True)
            
            message = client.messages.create(
                model=self.model,
                max_tokens=2000,
                temperature=0.3,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            
            summary = message.content[0].text.strip()
            print(f"✅ 文档分析完成，要点长度: {len(summary)} 字符", file=sys.stderr, flush=True)
            return summary
            
        except Exception as e:
            logger.error(f"文档分析失败: {str(e)}")
            print(f"❌ 文档分析失败: {str(e)}", file=sys.stderr, flush=True)
            raise Exception(f"文档分析失败: {str(e)}")

